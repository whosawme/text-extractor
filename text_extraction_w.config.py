import os
import pandas as pd
import docx
from docx import Document
import re
import xlsxwriter
# from xlwt  import Workbook

##---- CONFIGURATION VARIABLES ----##
Outputfilename="Exrtaction_outputV1.xlsx"
column1='doc_id'
column2='lesson_id'
column3='Headline'
column4='What'
column5='Based'

init_word_1= 'Headline'
stop_word_1= 'What'

start_word_2= 'What'
stop_word_2= 'Based'
conditional_word_2a= 'Lesson'
conditional_word_2b= 'Annex'
conditional_word_2c = init_word_1

start_word_3= 'Based'
stop_word_3= init_word_1
conditional_word_3a= 'Annex'

start_word_4= 'Lesson'
stop_word_4= init_word_1
conditional_word_4a= 'Annex'

##----- CONFIGURATION END -----##

strPath = os.getcwd()
translist = list()

### --- Directory Crawl --- ###
for file in os.listdir(strPath):
	if file.endswith('.docx'):
		doc_id = re.split('_|\s', os.path.splitext(file)[0])[0]
		print("_________\ncreated document_id: {} for: \n{}".format(doc_id,file))
		doc = docx.Document(file)
		sublesscounter=0
		# paralist = []
		for i in range(len(doc.paragraphs)):
			headline=''
			lesson_id=None
			what=''
			based=''
			if re.match(r'^%s\b'% init_word_1,doc.paragraphs[i].text): #.startswith('Headline') == True:
				sublesscounter+= 1
				lesson_id = str(doc_id+'-'+str(sublesscounter))
				
				print('_________\ninit.DEBUG:----> ',i)
				e=i
				o=i
				u=i

				while doc.paragraphs[e].text.startswith(stop_word_1)==False:
					headline = headline+' - '+doc.paragraphs[e].text
					print('H.DEBUG:----> ',e)
					e+=1
				if doc.paragraphs[e].text.startswith(start_word_2)==True:
					while doc.paragraphs[e].text.startswith(stop_word_2)==False and doc.paragraphs[e].text.startswith(conditional_word_2a)==False and doc.paragraphs[e].text.startswith(conditional_word_2b)==False and doc.paragraphs[e].text.startswith(conditional_word_2c)==False:
						what = what+' - ' +doc.paragraphs[e].text
						print('W.DEBUG:----> ',e)
						e+=1				
				if doc.paragraphs[e].text.startswith(start_word_3)==True:
					while doc.paragraphs[e].text.startswith(stop_word_3)==False and doc.paragraphs[e].text.startswith(conditional_word_3a)==False:
						print('B.DEBUG:----> ',e)
						based = based+'-'+doc.paragraphs[e].text
						e+=1
					if doc.paragraphs[e].text.startswith(start_word_3)==True:
						based = 'LESSON: '
						while doc.paragraphs[e].text.startswith(stop_word_3)==False and doc.paragraphs[e].text.startswith(conditional_word_4a)==False:
							print('L.DEBUG:----> ',e)
							based = based+'-'+doc.paragraphs[e].text
							e+=1
					
					# paradex = doc.paragraph[i].text.find('What')
					# tillwhat = doc.paragraphs[i:].find('What')

					# while i<tillwhat:
					# 	headline = headline+'-'+doc.paragraphs[i].text
					# 	i+=1
					# what = doc.paragraphs[tillwhat]					 


				print("--\n>>HEADLINE captured as: {}\n--\n>>WHAT captured as: {}\n--\n>>BASED captured as: {}\n".format(headline,what,based))
				sub_row_list = [doc_id,lesson_id,headline,what,based] #doc.paragraphs[i].text]
				translist.append(sub_row_list)
			else:
				pass
			
			# print("doc_id: {} \nlesson_id: {}\n paragraphs: \n".format(doc_id,lesson_id),doc.paragraphs[i].text)

		print("created lessons for {}\n\n--".format(file))

	else:
		pass
print("||Directory Crawl Complete")


#FROM LIST TO DATAFRAME
df=pd.DataFrame(translist, columns=[column1,column2,column3,column4,column5])

####----EXCEL WRITER----####

writer = pd.ExcelWriter(strPath+'\\'+Outputfilename)
df.to_excel(writer, sheet_name='Sheet1', startrow=0, startcol=0)
# workbook = xlsxwriter.Workbook('Testoutput.xlsx')
# workbook.autofilter('A1:D1')
writer.save()

print("||output document saved as '{} in {}'".format(Outputfilename,strPath))
print('||test complete')
